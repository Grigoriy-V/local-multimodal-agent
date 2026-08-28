"""Turning a document into text a model can be given, with its structure kept.

A document does not become model input the way a picture does. It is saved into
the person's workspace and read with a tool, because a fifty-page PDF pasted
into a turn would spend the context before the model had decided what mattered
in it.

What this module owes the model is the boundaries. "Page 7" and "Introduction"
are how a person refers to a place in a document, and an answer that cannot say
where it read something cannot be checked. So extraction returns labelled
sections rather than one string, and the tool that formats them keeps the
labels.

Every parser here is pointed at a file a stranger sent, so the format list is
short on purpose and the imports are local: a machine that never reads a PDF
does not need the library that reads one.
"""

from __future__ import annotations

import csv
import io
from dataclasses import dataclass

PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PLAIN = "text/plain"
MARKDOWN = "text/markdown"
CSV = "text/csv"

# What the assistant accepts, and what it calls each one when it says so.
DOCUMENT_MEDIA_TYPES = {
    PDF: "pdf",
    DOCX: "docx",
    PLAIN: "txt",
    MARKDOWN: "md",
    CSV: "csv",
}

# Telegram sends `text/plain` for a .md file and, for some clients, nothing at
# all. The suffix is the better evidence in both cases, so it wins when it names
# a format we know.
SUFFIX_MEDIA_TYPES = {
    ".pdf": PDF,
    ".docx": DOCX,
    ".txt": PLAIN,
    ".md": MARKDOWN,
    ".markdown": MARKDOWN,
    ".csv": CSV,
}

MAX_CSV_ROWS = 200


class DocumentError(ValueError):
    """A document cannot be read as the format it claims to be."""


@dataclass(frozen=True)
class Section:
    """One addressable piece of a document.

    `label` is what a person would say to point at it — "page 3", "Methods" —
    and is never invented: a format that has no such notion gets one section
    labelled after the whole document.
    """

    label: str
    text: str


def media_type_for(name: str, declared: str | None) -> str | None:
    """The format to read `name` as, or `None` if it is not a document.

    The suffix is consulted first because the declared type is whatever the
    sending client felt like saying.
    """

    suffix = name[name.rfind(".") :].lower() if "." in name else ""
    if suffix in SUFFIX_MEDIA_TYPES:
        return SUFFIX_MEDIA_TYPES[suffix]
    if declared in DOCUMENT_MEDIA_TYPES:
        return declared
    return None


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentError("the file is not text in any encoding this reads")


def _markdown_sections(text: str) -> list[Section]:
    """Split on ATX headings, which is the only structure Markdown really has."""

    sections: list[Section] = []
    label = "start"
    body: list[str] = []
    for line in text.splitlines():
        stripped = line.lstrip()
        if stripped.startswith("#") and stripped.lstrip("#").startswith(" "):
            if body:
                sections.append(Section(label, "\n".join(body).strip()))
                body = []
            label = stripped.lstrip("#").strip()
            continue
        body.append(line)
    if body or not sections:
        sections.append(Section(label, "\n".join(body).strip()))
    return [section for section in sections if section.text]


def _pdf_sections(data: bytes) -> list[Section]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = reader.pages
    except Exception as error:  # noqa: BLE001 - any parser failure is one answer
        raise DocumentError(f"the PDF could not be opened ({error})") from error
    sections = []
    for number, page in enumerate(pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 - one broken page is not a broken file
            text = ""
        sections.append(Section(f"page {number}", text.strip()))
    if not any(section.text for section in sections):
        # Saying this plainly matters more than it looks: a scanned PDF returns
        # empty pages, and an assistant that reports "the document is empty"
        # rather than "there is no text layer" sends the person away for good.
        raise DocumentError(
            f"the PDF has {len(sections)} page(s) and no text layer, so it is "
            "probably a scan; text can only be read from it by looking at it as "
            "an image"
        )
    return sections


def _docx_sections(data: bytes) -> list[Section]:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as error:  # noqa: BLE001 - any parser failure is one answer
        raise DocumentError(f"the .docx could not be opened ({error})") from error
    sections: list[Section] = []
    label = "start"
    body: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.style is not None and str(paragraph.style.name).startswith("Heading"):
            if body:
                sections.append(Section(label, "\n".join(body).strip()))
                body = []
            label = paragraph.text.strip() or label
            continue
        if paragraph.text.strip():
            body.append(paragraph.text)
    if body:
        sections.append(Section(label, "\n".join(body).strip()))
    # Tables come last rather than in place: python-docx walks paragraphs and
    # tables as two separate sequences, so their true interleaving is not
    # available here, and guessing at it would put a label on a lie.
    for number, table in enumerate(document.tables, start=1):
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
        ]
        if rows:
            sections.append(Section(f"table {number}", "\n".join(rows)))
    return [section for section in sections if section.text]


def _csv_sections(data: bytes) -> list[Section]:
    rows = list(csv.reader(io.StringIO(_decode(data))))
    if not rows:
        raise DocumentError("the CSV file has no rows")
    shown = rows[:MAX_CSV_ROWS]
    text = "\n".join(" | ".join(cell for cell in row) for row in shown)
    label = f"rows 1-{len(shown)} of {len(rows)}"
    return [Section(label, text)]


def read_sections(data: bytes, media_type: str) -> list[Section]:
    """Extract one document into labelled sections, or say why it cannot be."""

    if media_type == PDF:
        return _pdf_sections(data)
    if media_type == DOCX:
        return _docx_sections(data)
    if media_type == CSV:
        return _csv_sections(data)
    if media_type == MARKDOWN:
        return _markdown_sections(_decode(data))
    if media_type == PLAIN:
        text = _decode(data).strip()
        if not text:
            raise DocumentError("the file is empty")
        return [Section("whole file", text)]
    raise DocumentError(f"{media_type} is not a document format this can read")


# What a rendered page is sized to. The long side, in pixels: large enough that
# body text in a scan is legible to the model, small enough that a page is not
# an unreasonable share of one request. The model looks at the picture directly,
# so there is no OCR step for a bigger image to help.
PAGE_LONG_SIDE = 1400

# The serving limit is four images per prompt, and a turn may already carry the
# person's own photo and images from earlier messages. Two is what can be handed
# over without risking a refusal from the server.
MAX_PAGES_PER_VIEW = 2


def page_count(data: bytes) -> int:
    import pypdfium2 as pdfium

    try:
        document = pdfium.PdfDocument(io.BytesIO(data))
    except Exception as error:  # noqa: BLE001 - any parser failure is one answer
        raise DocumentError(f"the PDF could not be opened ({error})") from error
    try:
        return len(document)
    finally:
        document.close()


def render_pages(data: bytes, first: int, count: int) -> list[tuple[int, bytes]]:
    """Render pages `first..first+count-1` as PNGs, numbered from one.

    This is how a document with no text layer is read: the model is multimodal,
    so it looks at the page the way a person would, and nothing here recognises
    characters. That also means the failure mode is honest — an illegible scan
    looks illegible rather than becoming confident nonsense.
    """

    import pypdfium2 as pdfium

    try:
        document = pdfium.PdfDocument(io.BytesIO(data))
    except Exception as error:  # noqa: BLE001 - any parser failure is one answer
        raise DocumentError(f"the PDF could not be opened ({error})") from error
    # Closed explicitly: pypdfium2 holds native handles and complains at exit
    # about the ones it had to close itself, which would be noise in a log that
    # is already the only view into a deployed turn.
    try:
        total = len(document)
        if first < 1 or first > total:
            raise DocumentError(f"page {first} is outside a document of {total} page(s)")
        rendered: list[tuple[int, bytes]] = []
        for number in range(first, min(first + count, total + 1)):
            page = document[number - 1]
            try:
                longest = max(page.get_width(), page.get_height()) or 1
                scale = min(max(PAGE_LONG_SIDE / longest, 0.5), 4.0)
                image = page.render(scale=scale).to_pil()
            finally:
                page.close()
            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            rendered.append((number, buffer.getvalue()))
        return rendered
    finally:
        document.close()
