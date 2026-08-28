"""Documents keep their boundaries, and an unreadable one says why."""

from __future__ import annotations

import io
from pathlib import Path

import pytest

from app.attachments import AttachmentBytes, AttachmentError, admit_uploads, safe_filename
from app.documents import (
    CSV,
    MAX_PAGES_PER_VIEW,
    DOCX,
    MARKDOWN,
    PDF,
    PLAIN,
    DocumentError,
    media_type_for,
    read_sections,
)
from app.models import ToolCall
from app.tools import Toolbox
from app.tools.documents import document_tools


def a_pdf(pages: list[str]) -> bytes:
    """A real PDF with a text layer, built rather than fixtured."""

    from pypdf import PdfWriter

    try:  # reportlab is not a dependency; pypdf can only write blank pages.
        from reportlab.pdfgen.canvas import Canvas  # type: ignore
    except ImportError:
        pytest.skip("no PDF writer with a text layer is installed")
    buffer = io.BytesIO()
    canvas = Canvas(buffer)
    for text in pages:
        canvas.drawString(72, 720, text)
        canvas.showPage()
    canvas.save()
    writer = PdfWriter(clone_from=io.BytesIO(buffer.getvalue()))
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def a_docx(blocks: list[tuple[str, str]]) -> bytes:
    import docx

    document = docx.Document()
    for style, text in blocks:
        document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- what a file is read as --------------------------------------------------


def test_the_suffix_beats_whatever_the_client_claimed() -> None:
    """Telegram calls a .md file text/plain, and some clients say nothing.

    Reading a Markdown file as plain text loses exactly the thing a document
    reader is for, so the more specific evidence wins.
    """

    assert media_type_for("notes.md", "text/plain") == MARKDOWN
    assert media_type_for("notes.md", None) == MARKDOWN
    assert media_type_for("report.pdf", "application/octet-stream") == PDF


def test_something_that_is_not_a_document_says_so() -> None:
    assert media_type_for("photo.jpg", "image/jpeg") is None
    assert media_type_for("tool.exe", "application/x-msdownload") is None


# --- boundaries --------------------------------------------------------------


def test_a_pdf_keeps_its_page_numbers() -> None:
    sections = read_sections(a_pdf(["first page text", "second page text"]), PDF)

    assert [section.label for section in sections] == ["page 1", "page 2"]
    assert "first page" in sections[0].text


def test_markdown_keeps_its_headings() -> None:
    source = "intro line\n\n# Method\nhow it was done\n\n# Result\nwhat happened\n"

    sections = read_sections(source.encode("utf-8"), MARKDOWN)

    assert [section.label for section in sections] == ["start", "Method", "Result"]
    assert sections[2].text == "what happened"


def test_a_docx_keeps_its_headings() -> None:
    data = a_docx([("", "opening"), ("Heading 1", "Method"), ("", "how it was done")])

    sections = read_sections(data, DOCX)

    assert [section.label for section in sections] == ["start", "Method"]
    assert "how it was done" in sections[1].text


def test_a_csv_says_how_many_rows_it_did_not_show() -> None:
    data = ("a,b\n" + "1,2\n" * 500).encode("utf-8")

    sections = read_sections(data, CSV)

    assert sections[0].label == "rows 1-200 of 501"


def test_plain_text_in_a_windows_encoding_is_still_read() -> None:
    """A person on Windows sends cp1251, and that is not a corrupt file."""

    sections = read_sections("привет".encode("cp1251"), PLAIN)

    assert sections[0].text == "привет"


# --- failures that must be honest --------------------------------------------


def test_a_scanned_pdf_is_not_reported_as_an_empty_document() -> None:
    """The difference sends the person to a different next step.

    "The document is empty" makes them doubt the file. "There is no text layer"
    tells them it is a scan and that a picture of it is what to send.
    """

    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)

    with pytest.raises(DocumentError) as error:
        read_sections(buffer.getvalue(), PDF)

    assert "no text layer" in str(error.value)
    assert "scan" in str(error.value)


def test_a_broken_pdf_fails_as_a_document_not_as_a_crash() -> None:
    with pytest.raises(DocumentError):
        read_sections(b"not a pdf at all", PDF)


# --- the tool ----------------------------------------------------------------


async def test_the_tool_says_where_it_stopped_and_how_to_continue(tmp_path: Path) -> None:
    """A tool result goes straight into the next request unasked.

    So it is bounded, and it says what it did not show — otherwise the model
    answers from a fragment believing it has the whole thing.
    """

    body = "\n".join(f"# Section {index}\n{'text ' * 400}" for index in range(1, 12))
    (tmp_path / "long.md").write_text(body, encoding="utf-8")

    result = await Toolbox(document_tools(tmp_path)).run_async(
        ToolCall("d", "read_document", {"path": "long.md"})
    )
    text = result.content[0].text or ""

    assert "11 section(s)" in text
    assert "stopped after section" in text
    assert "from_section=" in text


async def test_the_tool_continues_from_where_it_stopped(tmp_path: Path) -> None:
    (tmp_path / "doc.md").write_text("# One\nalpha\n\n# Two\nbeta\n", encoding="utf-8")

    result = await Toolbox(document_tools(tmp_path)).run_async(
        ToolCall("d", "read_document", {"path": "doc.md", "from_section": 2})
    )
    text = result.content[0].text or ""

    assert "beta" in text
    assert "alpha" not in text


async def test_the_tool_refuses_a_path_outside_its_root(tmp_path: Path) -> None:
    root = tmp_path / "workspace"
    root.mkdir()
    (tmp_path / "outside.md").write_text("secret", encoding="utf-8")

    result = await Toolbox(document_tools(root)).run_async(
        ToolCall("d", "read_document", {"path": str(tmp_path / "outside.md")})
    )

    assert "outside the allowed root" in (result.content[0].text or "")


async def test_the_tool_names_what_it_can_read_when_asked_for_something_else(
    tmp_path: Path,
) -> None:
    (tmp_path / "photo.jpg").write_bytes(b"\xff\xd8\xff")

    result = await Toolbox(document_tools(tmp_path)).run_async(
        ToolCall("d", "read_document", {"path": "photo.jpg"})
    )
    text = result.content[0].text or ""

    assert "not a document" in text
    assert "pdf" in text


# --- admission ---------------------------------------------------------------


def test_a_sent_filename_cannot_climb_out_of_the_workspace(tmp_path: Path) -> None:
    """The name is written by whoever sent the file."""

    assert safe_filename("../../../etc/passwd") == "etc-passwd" or "/" not in safe_filename(
        "../../../etc/passwd"
    )
    assert ".." not in Path(safe_filename("..\\..\\windows\\system32\\x.txt")).parts


def test_two_files_with_one_name_are_two_files(tmp_path: Path) -> None:
    """Overwriting would answer the second question against a missing file."""

    uploads = [AttachmentBytes("report.pdf", PDF, b"first")]
    admit_uploads(uploads, tmp_path)
    admit_uploads([AttachmentBytes("report.pdf", PDF, b"second")], tmp_path)

    assert (tmp_path / "report.pdf").read_bytes() == b"first"
    assert (tmp_path / "report-2.pdf").read_bytes() == b"second"


def test_a_picture_still_goes_straight_to_the_model(tmp_path: Path) -> None:
    parts = admit_uploads([AttachmentBytes("photo.jpg", "image/jpeg", b"\xff\xd8")], tmp_path)

    assert [part.kind for part in parts] == ["image"]
    assert not list(tmp_path.iterdir())


def test_a_message_can_carry_both_a_picture_and_a_document(tmp_path: Path) -> None:
    parts = admit_uploads(
        [
            AttachmentBytes("photo.jpg", "image/jpeg", b"\xff\xd8"),
            AttachmentBytes("notes.md", "text/plain", b"# Notes\nbody"),
        ],
        tmp_path,
    )

    assert [part.kind for part in parts] == ["image", "text"]
    assert "notes.md" in (parts[1].text or "")
    assert (tmp_path / "notes.md").is_file()


def test_an_unreadable_format_never_reaches_the_workspace(tmp_path: Path) -> None:
    with pytest.raises(AttachmentError):
        admit_uploads([AttachmentBytes("tool.exe", "application/x-msdownload", b"MZ")], tmp_path)

    assert not list(tmp_path.iterdir())


# --- looking at a page -------------------------------------------------------


async def test_a_scan_is_read_by_looking_at_it(tmp_path: Path) -> None:
    """The point of the whole tool: no OCR step, no confident nonsense.

    The model is multimodal, so the page goes to it as a picture. What this
    asserts is that a real image comes back, in a media type the admission
    policy already accepts.
    """

    from app.attachments import MEDIA_KINDS

    (tmp_path / "scan.pdf").write_bytes(a_pdf(["page one", "page two"]))

    result = await Toolbox(document_tools(tmp_path)).run_async(
        ToolCall("v", "view_pages", {"path": "scan.pdf"})
    )

    kinds = [part.kind for part in result.content]
    assert kinds == ["text", "image"]
    image = result.content[1]
    assert image.media_type in MEDIA_KINDS
    assert image.data is not None and image.data.startswith(b"\x89PNG")
    assert "of 2" in (result.content[0].text or "")
    assert "page=2" in (result.content[0].text or "")


async def test_the_last_page_does_not_invite_a_next_one(tmp_path: Path) -> None:
    (tmp_path / "one.pdf").write_bytes(a_pdf(["only page"]))

    result = await Toolbox(document_tools(tmp_path)).run_async(
        ToolCall("v", "view_pages", {"path": "one.pdf", "page": 1})
    )

    assert "view_pages again" not in (result.content[0].text or "")


async def test_more_pages_than_the_server_accepts_are_never_returned(
    tmp_path: Path,
) -> None:
    """Four images per prompt is the serving limit, and a turn may already use some.

    Exceeding it is an HTTP 400, not a worse answer, so the cap is enforced here
    rather than hoped for in the schema.
    """

    (tmp_path / "many.pdf").write_bytes(a_pdf([f"page {index}" for index in range(1, 9)]))

    result = await Toolbox(document_tools(tmp_path)).run_async(
        ToolCall("v", "view_pages", {"path": "many.pdf", "page": 1, "pages": 8})
    )

    assert sum(1 for part in result.content if part.kind == "image") == MAX_PAGES_PER_VIEW


async def test_looking_past_the_end_says_how_long_the_document_is(tmp_path: Path) -> None:
    (tmp_path / "short.pdf").write_bytes(a_pdf(["only page"]))

    result = await Toolbox(document_tools(tmp_path)).run_async(
        ToolCall("v", "view_pages", {"path": "short.pdf", "page": 9})
    )

    assert "outside a document of 1 page(s)" in (result.content[0].text or "")


async def test_only_a_pdf_has_pages_to_look_at(tmp_path: Path) -> None:
    (tmp_path / "notes.md").write_text("# Notes\nbody", encoding="utf-8")

    result = await Toolbox(document_tools(tmp_path)).run_async(
        ToolCall("v", "view_pages", {"path": "notes.md"})
    )

    assert "not a PDF" in (result.content[0].text or "")


async def test_looking_at_a_page_refuses_a_path_outside_the_root(tmp_path: Path) -> None:
    root = tmp_path / "workspace"
    root.mkdir()
    (tmp_path / "outside.pdf").write_bytes(a_pdf(["secret"]))

    result = await Toolbox(document_tools(root)).run_async(
        ToolCall("v", "view_pages", {"path": str(tmp_path / "outside.pdf")})
    )

    assert "outside the allowed root" in (result.content[0].text or "")
